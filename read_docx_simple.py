
from docx import Document
import sys

def read_docx_simple(file_path):
    try:
        doc = Document(file_path)
        content = []
        
        content.append("--- PARAGRAPHS ---")
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text.strip())
        
        content.append("\n--- TABLES ---")
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # simplistic extraction, might duplicate text content if merged cells
                    # but good enough for reading
                    text = cell.text.strip().replace('\n', ' ')
                    if text:
                        row_data.append(text)
                if row_data:
                    content.append(" | ".join(row_data))
        
        return '\n'.join(content)

    except Exception as e:
        return f"Error: {e}"

if __name__ == "__main__":
    file_path = "d:/Year 4 semster 1/GRAD PROJECT/medtwin/NTRA-GP_Competition_template-2025_Ongoing-Academic-year-2025-2026 - UPDATED MEDTWIN.docx"
    print(read_docx_simple(file_path))
