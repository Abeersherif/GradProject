import docx
import re

# File Paths
input_file = "d:\\Year 4 semster 1\\GRAD PROJECT\\medtwin\\pdf_content.txt"
output_file = "d:\\Year 4 semster 1\\GRAD PROJECT\\medtwin\\MedTwin_Ai_Chronic_Update_Cleaned_v3.docx"

def sremove_emojis(text):
    # Regex for capturing emojis and symbols often used as emojis
    emoji_pattern = re.compile("[\U00010000-\U0010ffff]", flags=re.UNICODE)
    text = emoji_pattern.sub(r' ', text)
    # Also remove specific symbols seen in file
    text = text.replace('✅', '').replace('💊', '').replace('🩺', '').replace('🧠', '')
    return text.strip()

def clean_and_process():
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # 1. Remove Emojis
    content = sremove_emojis(content)

    # 2. Fix Broken/Split Agent Headers (Multi-pass regex)
    # Fix "Navigator + \n Safety Agent"
    content = re.sub(r'Navigator\s*\+\s*\n\s*Safety\s*Agent', 'Navigator + Safety Agent', content)
    
    # Fix "Simulation \n Agent" -> "Simulation Agent"
    # Applies to: Simulation, Coordinator, Monitoring
    content = re.sub(r'(Simulation|Coordinator|Monitoring)\s*\n\s*Agent', r'\1 Agent', content)
    
    # Fix "Emergency \n Triage Agent"
    content = re.sub(r'Emergency\s*\n\s*Triage\s*Agent', 'Emergency Triage Agent', content)

    # 3. Fix Merged Text "AgentConducts" -> "Agent\nConducts"
    # This splits the description to a new line so we can treat the Agent Name as a header
    content = re.sub(r'(Agent)(?=[A-Z])', r'\1\n', content)

    # 4. Remove Emergency Agent Block
    # Look for "Emergency Triage Agent" ... until "Simulation Agent"
    # We use re.DOTALL to match across lines
    content = re.sub(r'Emergency Triage Agent.*?Simulation Agent', 'Simulation Agent', content, flags=re.DOTALL)

    # 5. Create Word Doc
    doc = docx.Document()
    doc.add_heading("MedTwin – Updated Project Overview", 0)

    # Known Headers for formatting
    agent_names = [
        "Diagnostic Agent", "Navigator + Safety Agent", "Simulation Agent", 
        "Planner Agent", "Coordinator Agent", "Treatment Agent", 
        "Monitoring Agent", "Cognitive Agent"
    ]
    main_headers = ["The Big Idea", "The Problem We’re Solving", "What MedTwin Does", "The AI Helper Agents", "3D Visualization & Integration", "Overview", "How It Works"]

    lines = content.split('\n')
    
    for line in lines:
        line = strip_control_chars(line).strip()
        if not line:
            continue
            
        # Check Main Headers
        is_main = False
        for mh in main_headers:
            if mh in line and len(line) < len(mh) + 5:
                doc.add_heading(line, level=2)
                is_main = True
                break
        if is_main: continue

        # Check Agent Headers
        is_agent = False
        for ag in agent_names:
            # Exact match likely now that we split text
            if line == ag:
                doc.add_heading(line, level=3)
                is_agent = True
                break
        if is_agent: continue
        
        # Bullet points
        if line.startswith("•") or line.startswith("-") or re.match(r'^\d+\.', line):
            doc.add_paragraph(line, style='List Bullet')
            continue

        # Normal text
        doc.add_paragraph(line)

    doc.save(output_file)
    print(f"Created {output_file}")

def strip_control_chars(text):
    return "".join(ch for ch in text if ch.isprintable())


if __name__ == "__main__":
    clean_and_process()
