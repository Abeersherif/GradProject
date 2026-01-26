
from docx import Document
from docx.shared import Pt

def add_references(file_path):
    try:
        doc = Document(file_path)
        
        # Define references and their keywords
        references = {
            1: {
                "text": "International Diabetes Federation. IDF Diabetes Atlas, 10th edn. Brussels, Belgium: 2021. Available at: https://www.diabetesatlas.org",
                "keywords": ["537 million", "diabetes prevalence", "1 in 10 adults", "global health emergency"]
            },
            2: {
                "text": "American Diabetes Association. Standards of Medical Care in Diabetes—2024. Diabetes Care 2024;47(Supplement_1).",
                "keywords": ["cardiovascular disease", "chronic heart disease", "complications", "standards of care"]
            },
            3: {
                "text": "Elayan H, Aloqaily M, Guizani M. Digital Twins for Healthcare 4.0 - Recent Advances, Architecture, and Open Challenges. IEEE Access. 2021.",
                "keywords": ["digital twin", "virtual replica", "simulation", "healthcare 4.0", "predictive analysis"]
            },
            4: {
                "text": "Global Initiative for Chronic Obstructive Lung Disease (GOLD). Global Strategy for the Diagnosis, Management, and Prevention of Chronic Obstructive Pulmonary Disease. 2024 Report. Available at: https://goldcopd.org/",
                "keywords": ["COPD", "pulmonary", "lung disease"]
            }
        }

        # Helper to process text
        def process_text(text):
            original_text = text
            for ref_id, data in references.items():
                citation = f"[{ref_id}]"
                # Check keywords
                for keyword in data["keywords"]:
                    if keyword.lower() in text.lower() and citation not in text:
                        # Find the end of the sentence containing the keyword or just append to the keyword
                        # Simple approach: append to the keyword
                        idx = text.lower().find(keyword.lower())
                        if idx != -1:
                            # Insert citation after the keyword
                            insert_pos = idx + len(keyword)
                            text = text[:insert_pos] + " " + citation + text[insert_pos:]
            return text

        # 1. Iterate through Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                new_text = process_text(para.text)
                if new_text != para.text:
                    para.text = new_text

        # 2. Iterate through Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Avoid messing up the references section itself if it matches keywords
                    if "References" in cell.text: 
                        continue
                        
                    if cell.text.strip():
                        # naive text replacement for cells (might lose formatting, but safe for plain text cells)
                        new_text = process_text(cell.text)
                        if new_text != cell.text:
                            cell.text = new_text

        # 3. Add Bibliography
        # Find the "References" section or append to end
        found_ref_section = False
        for para in doc.paragraphs:
            if "References citation" in para.text or "References:" in para.text:
                found_ref_section = True
                # Move to next paragraphs to append
                # For simplicity in this script, we'll just append to the end of the document
                # if we can't easily find an insertion point inside the section.
                break
        
        doc.add_page_break()
        heading = doc.add_paragraph("References")
        heading.style = 'Heading 1'
        
        for ref_id, data in references.items():
            p = doc.add_paragraph(f"[{ref_id}] {data['text']}")
            p.style = 'Normal'

        doc.save(file_path)
        return "Successfully added references."

    except Exception as e:
        import traceback
        return f"Error: {e}\n{traceback.format_exc()}"

if __name__ == "__main__":
    file_path = "d:/Year 4 semster 1/GRAD PROJECT/medtwin/NTRA-GP_Competition_template-2025_Ongoing-Academic-year-2025-2026 - UPDATED MEDTWIN.docx"
    print(add_references(file_path))
