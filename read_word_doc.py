from docx import Document
import sys
import os

def read_docx(file_path):
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return

    try:
        doc = Document(file_path)
        print(f"--- Content of {file_path} ---")
        print("PARAGRAPHS:")
        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text)
        
        print("\nTABLES:")
        for i, table in enumerate(doc.tables):
            print(f"Table {i+1}:")
            for row in table.rows:
                row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                print(" | ".join(row_data))
            print("-" * 20)
            
    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        read_docx(sys.argv[1])
    else:
        print("Please provide a file path.")
