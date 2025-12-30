import docx
import os

# Paths
source_file = "NTRA-GP_Competition_template-2025_Ongoing-Academic-year-2025-2026 - UPDATED MEDTWIN.docx"
from docx.shared import Inches
output_file = "NTRA-GP_Competition_MedTwin_Filled_v9.docx"

# Project Data (Extracted from pdf_content.txt)
project_data = {
    "title": "MedTwin – Your Personal AI Hospital for Chronic Disease Care",
    "track": "Artificial Intelligence / Smart Healthcare",
    "university": "[University Name]",
    "department": "[Department/Faculty]",
    "industrial_partner": "Dell Technologies AI Empower Egypt",
    "partner_role": (
        "Support Type: Technical Mentorship & Capacity Building\n"
        "Dell Technologies AI Empower Egypt provides technical guidance on optimizing our AI models for deployment, "
        "access to potential computing resources for training our large-scale Digital Twin simulations, and "
        "mentorship on aligning the solution with industry standards for scalability and security."
    ),
    "modules_image": "system_architecture.png",
    
    "overview": (
        "Problem Definition:\n"
        "Chronic diseases are the leading cause of death and disability worldwide, with patients often suffering from multiple co-morbidities. "
        "Specifically, managing **COPD (Chronic Obstructive Pulmonary Disease), Type 2 Diabetes, and Chronic Heart Disease** requires a complex, delicate balance "
        "of medication, lifestyle adjustments, and continuous monitoring. Current healthcare systems are reactive; patients drift until an acute crisis occurs. "
        "Doctors lack real-time visibility into the daily progression of these conditions, leading to delayed interventions and avoidable hospitalizations.\n\n"
        "Approach & Tools (The MedTwin Solution):\n"
        "MedTwin is a comprehensive AI-powered **Digital Twin Ecosystem** designed to shift care from reactive to proactive. "
        "It utilizes a 'Collaborative Mode' Multi-Agent System where specialized agents handle every phase of care:\n"
        "1. **Diagnostic Agent**: Conducts medically approved interviews to analyze initial symptoms and update the twin.\n"
        "2. **Navigator & Safety Agent**: Plans next steps (tests) and checks for medication conflicts/allergies.\n"
        "3. **Simulation Agent**: Virtually tests treatment outcomes (Projected success rates) on the patient's Digital Twin.\n"
        "4. **Planner Agent**: Drafts detailed care plans integrating simulation data for doctor review.\n"
        "5. **Coordinator Agent**: Generates reports and synchronizes data across the entire care team (Doctors, Nurses).\n"
        "6. **Treatment Agent**: Handles daily medication tracking, reminders, and adherence feedback.\n"
        "7. **Monitoring Agent**: Continuously observes vitals from wearables to detect deviations.\n"
        "8. **Cognitive Agent**: Focuses on prevention by predicting long-term risks (e.g., 5-year heart failure risk).\n\n"
        "(iii) **Overview of System Modules**\n"
        "The MedTwin architecture is composed of three interconnected layers (see diagram below):\n"
        "1. **User Interface Layer**:\n"
        "   - *Patient Mobile App*: A React Native/Flutter application providing the chat interface, medication reminders, and 3D visualization.\n"
        "   - *Doctor Web Dashboard*: A secure portal for clinicians to validate AI decisions and monitor patient cohorts.\n"
        "2. **Intelligent Core (The Brain)**:\n"
        "   - *Multi-Agent Orchestrator*: A Python-based backend where the 8 specialized agents collaborate.\n"
        "   - *Digital Twin Engine*: The simulation module that updates the patient’s virtual state based on incoming data streams.\n"
        "3. **Data Layer**:\n"
        "   - *Real-time Database*: Stores time-series data from wearables (heart rate, glucose).\n"
        "   - *Knowledge Graph*: Contains medical protocols for COPD, Diabetes, and Heart Failure management."
    ),
    
    "impact": (
        "1. **Clinical Impact (Patient Outcome Improvements)**:\n"
        "   - **Reduced Readmissions**: By continuously monitoring patients with Heart Failure and COPD, we can detect decompensation early, potentially reducing hospital readmission rates by 20-30%.\n"
        "   - **Personalized Precision**: Type 2 Diabetes management is tailored to the individual's live biology, not just general guidelines, improving HbA1c control.\n\n"
        "2. **Economic Impact (Cost Efficiency)**:\n"
        "   - **Healthcare Savings**: Preventing complex complications (like diabetic ketoacidosis or heart attacks) saves immense costs for the healthcare system (state or insurance).\n"
        "   - **Optimized Resources**: Automating routine follow-ups allows doctors to see more patients in less time, increasing the throughput of public clinics.\n\n"
        "3. **Social Impact (Accessibility & Quality of Life)**:\n"
        "   - **Equity**: Brings 'Consultant-level' monitoring to patients in remote areas where specialists are scarce.\n"
        "   - **Patient Empowerment**: Use of simple, AI-driven daily advice reduces the mental burden of self-care for elderly patients and their families.\n\n"
        "4. **Scientific & Strategic Impact**:\n"
        "   - **Advancing Digital Health**: Establishes a framework for 'Digital Twins' in Egypt, moving beyond simple EMRs (Electronic Medical Records) to active biological simulation.\n"
        "   - **Data Ecosystem**: Creates a valuable, anonymized longitudinal dataset for future research into multi-morbidity progression in the Egyptian population."
    ),
    
    "novelty": (
        "Scientific & Technical Novelty:\n"
        "While many apps track single conditions (e.g., just diabetes), MedTwin is built to handle **multi-morbidity**. \n"
        "- **True Digital Twin**: We don't just log data; we simulate physiological responses. Our Simulation Agent can predict risk trajectories years into the future.\n"
        "- **Multi-Agent Orchestration**: Unlike simple chatbots, our system uses distinct agents (Planner, Monitor, Critic) that debate and collaborate to find the optimal care plan, mimicking a medical board.\n\n"
        "Competitive Features:\n"
        "- **vs. Ada/Babylon**: These are primary care triage tools. MedTwin provides *continuous* long-term management specific to complex chronic needs.\n"
        "- **vs. MySugr/Livongo**: These focus on single diseases. MedTwin integrates data for patients who might have both Diabetes and Heart Failure, checking for conflicting contraindications in real-time.\n"
        "- **3D Visualization**: Includes an interactive 3D anatomical model that helps patients visualize 'organ stress', improving psychological adherence to therapy."
    ),
    
    "deliverables": (
        "1. **Patient Mobile Application**: A user-friendly interface for symptom reporting (Voice/Text), medication adherence tracking, and receiving AI-guided lifestyle alerts.\n"
        "2. **Healthcare Provider Dashboard**: A web portal for doctors to view patient 'Twins', approve AI-suggested care plans, and monitor cohort risks.\n"
        "3. **Intelligent Backend Core**: The server-side Python environment hosting the Multi-Agent System (using LangGraph/LangChain) and the Digital Twin state manager.\n"
        "4. **3D Interactive Module**: A Web-based visualization component showing real-time organ status (Heart, Lungs, Pancreas) based on live biometrics.\n"
        "5. **Validation Study**: A pilot report comparing MedTwin's risk flags against historical patient data."
    ),
    
    "business_plan": (
        "Market Needs:\n"
        "The global digital health market is exploding, but there is a lack of integrated tools for complex chronic patients. "
        "Hospitals are penalized for high readmission rates (especially for Heart Failure and COPD), creating a direct B2B incentive for our solution.\n\n"
        "Target Audience:\n"
        "- Primary: Hospitals and Private Clinics managing large pools of chronic patients.\n"
        "- Secondary: Insurance companies seeking to reduce long-term claim costs.\n"
        "- User Base: Patients diagnosed with Type 2 Diabetes, COPD, or Cardiovascular disease.\n\n"
        "Revenue Model:\n"
        "1. **B2B SaaS**: Licensing the Doctor Dashboard and Triage system to hospitals (Monthly Recurring Revenue per patient).\n"
        "2. **Freemium B2C**: Basic app is free; 'Premium Twin' features (advanced simulations, 3D views) for individual users.\n\n"
        "Growth Strategy:\n"
        "Phase 1: Pilot with local university hospitals.\n"
        "Phase 2: Expansion to regional private clinics.\n"
        "Phase 3: Integration with national insurance providers."
    ),
    
    "expenses": [
        ("Cloud Hosting (AWS/GCP)", "Infrastructure", "Server costs for AI agents & DB", "5000"),
        ("MedTwin Domain & SSL", "Software/Web", "Domain registration and security", "1000"),
        ("LLM API Credits", "Software/Service", "OpenAI/Google Gemini API usage", "3000"),
        ("3D Assets & Licenses", "Assets", "Unity/Three.js models", "2000"),
        ("User Testing Incentives", "Marketing/Research", "Participant compensation", "1000")
    ]
}

def fill_form():
    if not os.path.exists(source_file):
        print(f"Error: {source_file} not found.")
        return

    doc = docx.Document(source_file)
    
    # --- FILLING TABLES BASED ON INSPECTION ---
    
    # Table 0: Project Info
    # Row 0, Cell 1: Project Title
    doc.tables[0].rows[0].cells[1].text = project_data["title"]
    # Row 1, Cell 1: Project Track
    doc.tables[0].rows[1].cells[1].text = project_data["track"]
    # Row 2, Cell 1: University
    doc.tables[0].rows[2].cells[1].text = project_data["university"]
    # Row 3, Cell 1: Department
    doc.tables[0].rows[3].cells[1].text = project_data["department"]
    # Row 4, Cell 1: Industrial Partner
    doc.tables[0].rows[4].cells[1].text = project_data["industrial_partner"]
    
    # Table 3: Overview
    # Row 2, Cell 0: (Merged cell usually) - Fill the empty row below header
    # Inspect showed Row 2 is empty
    cell = doc.tables[3].rows[2].cells[0]
    cell.text = project_data["overview"]
    
    # Add Image
    # (User opted for manual diagram creation, so we skip programmatic image insertion here)
    # You can paste the diagram manually into the docx.
    
    # Table 4: Impact
    # Row 2, Cell 0
    doc.tables[4].rows[2].cells[0].text = project_data["impact"]
    
    # Table 5: Novelty and Deliverables
    # Row 2, Cell 0: Novelty
    doc.tables[5].rows[2].cells[0].text = project_data["novelty"]
    # Row 5, Cell 0: Deliverables
    doc.tables[5].rows[5].cells[0].text = project_data["deliverables"]
    
    # Table 6: Business Plan & Partner Role
    # Row 2, Cell 0: Market Analysis
    doc.tables[6].rows[2].cells[0].text = project_data["business_plan"]
    # Row 5, Cell 0: Partner Role (Table 6, Row 3 is Header "ROLE OF...", Row 4 is Question, Row 5 is Answer space)
    doc.tables[6].rows[5].cells[0].text = project_data["partner_role"]
    
    # Table 7: Expenses
    # Header is Row 0. Items start at Row 1.
    # Columns: Item, Type, Specifications, Cost
    # We have 5 items in our list.
    expenses = project_data["expenses"]
    table_expenses = doc.tables[7]
    
    total_cost = 0
    
    # The table has many empty rows; we fill them.
    # Start checking from Row 1
    current_row_idx = 1
    for item in expenses:
        if current_row_idx >= len(table_expenses.rows):
            break # No more rows
        
        row = table_expenses.rows[current_row_idx]
        # Make sure it's not the Total row (last one usually)
        if "Total Project Cost" in row.cells[0].text:
             break
             
        row.cells[0].text = item[0] # Item
        row.cells[1].text = item[1] # Type
        row.cells[2].text = item[2] # Specs
        row.cells[3].text = item[3] # Cost
        
        try:
            total_cost += int(item[3])
        except:
            pass
            
        current_row_idx += 1
        
    # Find Total Row
    # Inspect says "Table 7... Row 11: ['Total Project Cost', 'Total Project Cost', '', '']"
    # It might vary, so let's search for it just in case
    for row in table_expenses.rows:
        if "Total Project Cost" in row.cells[0].text or "Total Project Cost" in row.cells[1].text:
            # Usually the last cell has the value or the cell next to the label
            # Based on inspection: ['Total Project Cost', 'Total Project Cost', '', '']
            # It seems merged or duplicated text. We'll try to put it in the last cell
            row.cells[-1].text = str(total_cost)
            break

    doc.save(output_file)
    print(f"Successfully created {output_file}")

if __name__ == "__main__":
    fill_form()
