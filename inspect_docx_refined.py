import docx

file_path = "d:\\Year 4 semster 1\\GRAD PROJECT\\medtwin\\MedTwin_Ai_Chronic_Update_Cleaned.docx"

def print_structure():
    doc = docx.Document(file_path)
    print("--- DOCUMENT STRUCTURE ---")
    
    # Iterate through all block-level elements (paragraphs and tables) strictly in order? 
    # python-docx separates them. We'll just list tables with enough context.
    
    print(f"Total Tables: {len(doc.tables)}")
    
    for i, table in enumerate(doc.tables):
        print(f"\n[TABLE {i}]")
        # Print first few rows to identify
        for r_idx, row in enumerate(table.rows):
            cells = [c.text.strip().replace('\n', ' ')[:50] for c in row.cells]
            print(f"  Row {r_idx}: {cells}")
            if r_idx > 5: # Limit output
                print("  ... (more rows)")
                break

if __name__ == "__main__":
    print_structure()
